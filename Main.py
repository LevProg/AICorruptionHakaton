from tkinter import *
import tkinter.font as tkFont
from tkinter import filedialog as fd
import threading
import docx
import pickle
import sys
import os
import comtypes.client
from catboost import CatBoostClassifier, Pool
from docx2pdf import convert
from docx.enum.text import WD_COLOR_INDEX
from tkdocviewer import *

best_fac = ['3_9', '3_5', '4_3']

model=CatBoostClassifier()
models = pickle.load(open('models.pickle', 'rb'))
TRESHOLD = 0.75


def Classificate(filenam):
    global allCntOfCorruption
    allCntOfCorruption=0
    filename=filenam
    Predict()

def Predict():
    doc = docx.Document(filename)
    cnt=0
    for paragraph in doc.paragraphs:
        indicator = False
        for fac in best_fac:  
            par = paragraph.text
            vectorizer, clf = models[fac]
            if clf.predict_proba(vectorizer.transform([par]))[0][1]>TRESHOLD:
                cnt+=1
                indicator = True
                break
        if indicator:
            for run in paragraph.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW#Highlight the paragraph with the found corruption factor

    allCntOfCorruption=cnt
    doc.save('new.docx')

    wdFormatPDF = 17
    in_file = os.path.abspath('new.docx')
    out_file = os.path.abspath('new.pdf')
    word = comtypes.client.CreateObject('Word.Application')#docx to PDF
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    v.display_file("new.pdf")

    predictLabel.configure(text="Акт содержит: " + str(allCntOfCorruption) + " нарушени(е/й)", bg=_from_rgb((49,65,78)))
    btn.pack(padx=30, pady=10,side='bottom',fill="x")
    predictLabel.pack(padx=30,pady=20,fill="both")

def _from_rgb(rgb):
    return "#%02x%02x%02x" % rgb#translates an rgb tuple of int to a tkinter friendly color code

def choose_file():
    global filename
    threadClassificate = threading.Thread(target=Predict)#создание потока для классификации
    filetypes = (("Изображение", "*.doc *.docx"),
            ("Текстовый файл", "*.txt"),
            ("Любой", "*"))
    filename = fd.askopenfilename(title="Открыть файл", initialdir="/",
                                filetypes=filetypes)
    if filename:
        threadClassificate.run()
        
def save_file():
    filename = fd.asksaveasfile(initialfile = 'new.docx',
    defaultextension=".docx",filetypes=[("All Files","*.*"),("Text Documents","*.txt")])
    print(filename.name)
    doc=docx.Document('new.docx')
    doc.save(filename.name)


def window():
    root = Tk()
    root["bg"] = _from_rgb((36,52,65))
    root.geometry('1080x1080')
    root.title("Детектор")
    mainFont = tkFont.Font(family="Impact", size=16)

    global predictLabel, predictImage, btn, v
    predictLabel = Label(root, text='',font=mainFont, fg=_from_rgb((41,221,209)),bg=_from_rgb((36,52,65)))
    choseLabel = Label(root, text='Выберите документ для проверки на коррупционные факторы', fg=_from_rgb((41,221,209)),bg=_from_rgb((49,65,78)),font=mainFont)
    btn_file = Button(text="Выбор файла",font=mainFont,command=choose_file,fg=_from_rgb((41,221,209)),background=_from_rgb((49,65,78)),height=3,width=20)
    btn= Button(root, font=mainFont,command= lambda:save_file(),fg=_from_rgb((41,221,209)),background=_from_rgb((49,65,78)), text= "Сохранить файл",height=3,width=20)
    v = DocViewer(root)
    lf = Frame(root)

    v.pack(expand=1, fill="both",padx=50,pady=20,side='left')
    choseLabel.pack(pady=10,
                    ipadx=10,
                    ipady=10,
                    fill='y')
    btn_file.pack(padx=30, pady=10,side='top',fill="x")



    x = root.winfo_screenwidth()  # размер  по горизонтали
    y = root.winfo_screenheight()  # размер по вертикали
    root.geometry('{}x{}'.format(int(x), int(y)))
    root.mainloop()



windowThread = threading.Thread(target=window)
threadClassificate = threading.Thread(target=Predict)#создание потока для классификации
windowThread.start()


